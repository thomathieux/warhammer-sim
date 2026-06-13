# -*- coding: utf-8 -*-
"""
Génère un rapport Word sur les représentations statistiques de probabilités.
Usage : python tools/gen_stat_report.py
Output : rapport_stats_proba.docx (à la racine du projet)
"""

import io
import numpy as np
import matplotlib
matplotlib.use("Agg")
import matplotlib.pyplot as plt
import matplotlib.patches as mpatches
from matplotlib.patches import FancyArrowPatch
import matplotlib.gridspec as gridspec
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import scipy.stats as stats

# ---------------------------------------------------------------------------
# Palette DA (cohérente avec l'app)
# ---------------------------------------------------------------------------
C_BG        = "#F5F0E8"
C_BG2       = "#FAF7F2"
C_INK       = "#1a1a2e"
C_GRID      = "#D8D0C0"
C_TERRA     = "#C47A5A"
C_BLUE      = "#7A9EC4"
C_GREEN     = "#7A9E7E"
C_ROSE      = "#C4867A"
C_SAND      = "#C4A87A"
C_LIGHT     = "#E8F0EC"

plt.rcParams.update({
    "figure.facecolor": C_BG,
    "axes.facecolor":   C_BG2,
    "axes.edgecolor":   C_INK,
    "axes.labelcolor":  C_INK,
    "xtick.color":      C_INK,
    "ytick.color":      C_INK,
    "grid.color":       C_GRID,
    "grid.linewidth":   0.8,
    "text.color":       C_INK,
    "font.family":      "serif",
    "font.size":        11,
    "axes.titlesize":   13,
    "axes.titleweight": "bold",
})

rng = np.random.default_rng(42)

# ---------------------------------------------------------------------------
# Données simulées — distribution WH40K réaliste
# ---------------------------------------------------------------------------
def sim_damage(n=2000):
    """
    Simule des dégâts sur n rounds.
    Mélange : 20% de rounds à 0 dégât (toutes sauvegardes réussies)
    + distribution Binomiale/Poisson pour le reste.
    """
    zeros = rng.random(n) < 0.18
    base = rng.poisson(lam=5.5, size=n).astype(float)
    base[zeros] = 0
    return np.clip(base, 0, 24).astype(int)


def sim_kills(dmg, wounds_per_model=3, n_models=6):
    return np.minimum(dmg // wounds_per_model, n_models)


dmg = sim_damage(2000)
kills = sim_kills(dmg)
dmg_weapon2 = sim_damage(2000) + rng.poisson(2, 2000)  # arme alternative
dmg_weapon2 = np.clip(dmg_weapon2, 0, 28).astype(int)

# Phase data pour funnel
phases = {
    "Attaques":    20,
    "Touches":     13.2,
    "Blessures":   8.8,
    "Non sauveg.": 5.5,
    "Dégâts":      9.8,
}

# ---------------------------------------------------------------------------
# Helpers figure
# ---------------------------------------------------------------------------
def fig_to_bytes(fig, dpi=150):
    buf = io.BytesIO()
    fig.savefig(buf, format="png", dpi=dpi, bbox_inches="tight")
    buf.seek(0)
    plt.close(fig)
    return buf


def styled_ax(ax, title="", xlabel="", ylabel=""):
    ax.set_title(title, pad=8)
    ax.set_xlabel(xlabel)
    ax.set_ylabel(ylabel)
    ax.grid(True, alpha=0.5)
    ax.spines["top"].set_visible(False)
    ax.spines["right"].set_visible(False)
    return ax


# ---------------------------------------------------------------------------
# 1. Histogramme
# ---------------------------------------------------------------------------
def chart_histogram():
    fig, ax = plt.subplots(figsize=(6, 3.5))
    n_bins = max(dmg) - min(dmg) + 1
    ax.hist(dmg, bins=n_bins, color=C_TERRA, edgecolor=C_INK, linewidth=0.6, rwidth=0.85)
    ax.axvline(np.mean(dmg), color=C_INK, linestyle="--", linewidth=1.4, label=f"Moy. {np.mean(dmg):.1f}")
    ax.legend(framealpha=0.7)
    styled_ax(ax, "Histogramme — Distribution des dégâts", "Dégâts", "Fréquence")
    return fig_to_bytes(fig)


# ---------------------------------------------------------------------------
# 2. KDE (Kernel Density Estimate)
# ---------------------------------------------------------------------------
def chart_kde():
    fig, ax = plt.subplots(figsize=(6, 3.5))
    kde = stats.gaussian_kde(dmg, bw_method=0.4)
    x = np.linspace(0, max(dmg) + 2, 300)
    y = kde(x)
    ax.fill_between(x, y, alpha=0.35, color=C_TERRA)
    ax.plot(x, y, color=C_TERRA, linewidth=2.5)
    ax.axvline(np.mean(dmg), color=C_INK, linestyle="--", linewidth=1.3, label=f"Moy. {np.mean(dmg):.1f}")
    ax.legend(framealpha=0.7)
    styled_ax(ax, "KDE — Densité de probabilité lissée", "Dégâts", "Densité")
    return fig_to_bytes(fig)


# ---------------------------------------------------------------------------
# 3. Histogramme + KDE superposés
# ---------------------------------------------------------------------------
def chart_hist_kde():
    fig, ax = plt.subplots(figsize=(6, 3.5))
    n_bins = max(dmg) - min(dmg) + 1
    ax.hist(dmg, bins=n_bins, color=C_TERRA, edgecolor=C_INK, linewidth=0.5,
            rwidth=0.85, density=True, alpha=0.65, label="Histogramme (densité)")
    kde = stats.gaussian_kde(dmg, bw_method=0.4)
    x = np.linspace(0, max(dmg) + 2, 300)
    ax.plot(x, kde(x), color=C_INK, linewidth=2.2, label="KDE")
    ax.legend(framealpha=0.7)
    styled_ax(ax, "Histogramme + KDE superposés", "Dégâts", "Densité")
    return fig_to_bytes(fig)


# ---------------------------------------------------------------------------
# 4. CDF (Fonction de répartition)
# ---------------------------------------------------------------------------
def chart_cdf():
    fig, ax = plt.subplots(figsize=(6, 3.5))
    sorted_d = np.sort(dmg)
    p = np.arange(1, len(sorted_d) + 1) / len(sorted_d)
    ax.step(sorted_d, p * 100, color=C_GREEN, linewidth=2.5, where="post")
    ax.axhline(50, color=C_INK, linestyle=":", linewidth=1, alpha=0.6)
    ax.text(max(dmg) * 0.85, 52, "P50", fontsize=9, color=C_INK, alpha=0.8)
    ax.set_ylim(0, 105)
    styled_ax(ax, "CDF — P(dégâts ≤ X)", "Dégâts", "Probabilité (%)")
    return fig_to_bytes(fig)


# ---------------------------------------------------------------------------
# 5. Survival function / CCDF — P(X >= x)
# ---------------------------------------------------------------------------
def chart_survival():
    fig, ax = plt.subplots(figsize=(6, 3.5))
    x_vals = np.arange(0, max(dmg) + 2)
    surv = [np.mean(dmg >= v) * 100 for v in x_vals]
    ax.fill_between(x_vals, surv, alpha=0.25, color=C_GREEN, step="post")
    ax.step(x_vals, surv, color=C_GREEN, linewidth=2.5, where="post")
    # Marquer P50 (médiane)
    med = np.median(dmg)
    ax.axvline(med, color=C_TERRA, linestyle="--", linewidth=1.4, label=f"Médiane ({med:.0f})")
    ax.legend(framealpha=0.7)
    ax.set_ylim(0, 105)
    styled_ax(ax, "Survival function — P(dégâts ≥ X)", "Dégâts", "Probabilité (%)")
    return fig_to_bytes(fig)


# ---------------------------------------------------------------------------
# 6. Box plot
# ---------------------------------------------------------------------------
def chart_boxplot():
    fig, ax = plt.subplots(figsize=(6, 3.5))
    bp = ax.boxplot(
        dmg, vert=False, patch_artist=True, widths=0.5,
        medianprops=dict(color=C_INK, linewidth=2),
        boxprops=dict(facecolor=C_TERRA, alpha=0.7, edgecolor=C_INK),
        whiskerprops=dict(color=C_INK, linewidth=1.2),
        capprops=dict(color=C_INK, linewidth=1.5),
        flierprops=dict(marker="o", markerfacecolor=C_TERRA, markersize=4, alpha=0.4),
    )
    ax.set_yticks([])
    p25, p50, p75 = np.percentile(dmg, [25, 50, 75])
    ax.text(p25, 1.32, f"P25\n{p25:.0f}", ha="center", fontsize=8, color=C_INK)
    ax.text(p50, 1.32, f"P50\n{p50:.0f}", ha="center", fontsize=8, color=C_INK)
    ax.text(p75, 1.32, f"P75\n{p75:.0f}", ha="center", fontsize=8, color=C_INK)
    styled_ax(ax, "Box plot — Résumé 5-nombres", "Dégâts", "")
    return fig_to_bytes(fig)


# ---------------------------------------------------------------------------
# 7. Violin plot
# ---------------------------------------------------------------------------
def chart_violin():
    fig, ax = plt.subplots(figsize=(6, 3.5))
    parts = ax.violinplot(dmg, vert=False, showmedians=True, showextrema=True)
    for pc in parts["bodies"]:
        pc.set_facecolor(C_TERRA)
        pc.set_edgecolor(C_INK)
        pc.set_alpha(0.7)
    for part in ["cmedians", "cmaxes", "cmins", "cbars"]:
        parts[part].set_color(C_INK)
        parts[part].set_linewidth(1.5)
    ax.set_yticks([])
    styled_ax(ax, "Violin plot — Distribution complète + densité", "Dégâts", "")
    return fig_to_bytes(fig)


# ---------------------------------------------------------------------------
# 8. Strip plot (jitter)
# ---------------------------------------------------------------------------
def chart_strip():
    fig, ax = plt.subplots(figsize=(6, 3.5))
    sample = rng.choice(dmg, size=300, replace=False)
    jitter = rng.uniform(-0.3, 0.3, size=len(sample))
    ax.scatter(sample, jitter, alpha=0.35, color=C_TERRA, s=18, edgecolors="none")
    ax.axvline(np.mean(dmg), color=C_INK, linestyle="--", linewidth=1.5, label=f"Moy. {np.mean(dmg):.1f}")
    ax.set_yticks([])
    ax.set_ylim(-0.6, 0.6)
    ax.legend(framealpha=0.7)
    styled_ax(ax, "Strip plot — Points individuels avec jitter", "Dégâts", "")
    return fig_to_bytes(fig)


# ---------------------------------------------------------------------------
# 9. Raincloud plot (violin + strip + box combinés)
# ---------------------------------------------------------------------------
def chart_raincloud():
    fig, ax = plt.subplots(figsize=(6, 4))
    # Violin (moitié haute)
    parts = ax.violinplot(dmg, vert=False, positions=[0.5], showmedians=False, showextrema=False)
    for pc in parts["bodies"]:
        pc.set_facecolor(C_TERRA)
        pc.set_edgecolor(C_INK)
        pc.set_alpha(0.65)
        # Couper le violin en deux (ne garder que la moitié haute)
        m = pc.get_paths()[0]
        verts = m.vertices
        verts[verts[:, 1] < 0.5, 1] = 0.5
    # Box
    bp = ax.boxplot(dmg, vert=False, positions=[0.5], widths=0.12, patch_artist=True,
                    medianprops=dict(color=C_INK, linewidth=2),
                    boxprops=dict(facecolor=C_BG2, edgecolor=C_INK, alpha=0.9),
                    whiskerprops=dict(color=C_INK), capprops=dict(color=C_INK),
                    flierprops=dict(marker="", markersize=0))
    # Jitter (pluie sous le violin)
    sample = rng.choice(dmg, size=300, replace=False)
    jitter = rng.uniform(0.05, 0.45, size=len(sample))
    ax.scatter(sample, jitter, alpha=0.25, color=C_BLUE, s=12, edgecolors="none")
    ax.set_yticks([])
    ax.set_ylim(-0.1, 1.1)
    styled_ax(ax, "Raincloud plot — Cloud + Box + Pluie", "Dégâts", "")
    return fig_to_bytes(fig)


# ---------------------------------------------------------------------------
# 10. Percentile bands (bandes de confiance)
# ---------------------------------------------------------------------------
def chart_percentile_bands():
    fig, ax = plt.subplots(figsize=(6, 3.5))
    x_vals = np.arange(0, max(dmg) + 2)
    surv = np.array([np.mean(dmg >= v) * 100 for v in x_vals])
    ax.fill_between(x_vals, 0, surv, alpha=0.08, color=C_GREEN, step="post", label="Plage totale")
    # P10 - P90
    p10 = np.percentile(dmg, 10)
    p90 = np.percentile(dmg, 90)
    mask = (x_vals >= p10) & (x_vals <= p90)
    ax.fill_between(x_vals[mask], 0, surv[mask], alpha=0.3, color=C_GREEN, step="post", label="P10–P90 (80%)")
    # P25 - P75
    p25 = np.percentile(dmg, 25)
    p75 = np.percentile(dmg, 75)
    mask2 = (x_vals >= p25) & (x_vals <= p75)
    ax.fill_between(x_vals[mask2], 0, surv[mask2], alpha=0.5, color=C_GREEN, step="post", label="P25–P75 (50%)")
    ax.step(x_vals, surv, color=C_GREEN, linewidth=2, where="post")
    ax.axvline(np.mean(dmg), color=C_TERRA, linestyle="--", linewidth=1.5,
               label=f"Moy. {np.mean(dmg):.1f}")
    ax.set_ylim(0, 105)
    ax.legend(fontsize=9, framealpha=0.7, loc="upper right")
    styled_ax(ax, "Bandes de percentiles — Intervalles de confiance", "Dégâts", "P(≥ X) %")
    return fig_to_bytes(fig)


# ---------------------------------------------------------------------------
# 11. Funnel chart (phases de combat)
# ---------------------------------------------------------------------------
def chart_funnel():
    fig, ax = plt.subplots(figsize=(6, 4))
    labels = list(phases.keys())
    values = list(phases.values())
    colors = [C_TERRA, C_SAND, C_GREEN, C_BLUE, C_ROSE]
    y_pos = np.arange(len(labels))
    max_v = max(values)
    bars = ax.barh(y_pos, values, align="center", height=0.6,
                   color=colors, edgecolor=C_INK, linewidth=0.7)
    for i, (v, bar) in enumerate(zip(values, bars)):
        pct = v / values[0] * 100
        ax.text(v + 0.3, i, f"{v:.1f}  ({pct:.0f}%)", va="center", fontsize=10)
    ax.set_yticks(y_pos)
    ax.set_yticklabels(labels, fontsize=11)
    ax.set_xlim(0, max_v * 1.45)
    ax.invert_yaxis()
    ax.set_xlabel("Valeur moyenne")
    ax.grid(True, axis="x", alpha=0.4)
    ax.spines["top"].set_visible(False)
    ax.spines["right"].set_visible(False)
    ax.set_title("Funnel chart — Phases de combat", pad=8, fontweight="bold")
    return fig_to_bytes(fig)


# ---------------------------------------------------------------------------
# 12. Gauge (jauge de menace)
# ---------------------------------------------------------------------------
def chart_gauge():
    fig, ax = plt.subplots(figsize=(6, 3.5), subplot_kw={"projection": "polar"})
    score = 0.63  # 63% de menace

    # Fond arc
    theta_max = np.pi
    theta = np.linspace(np.pi, 0, 300)
    zone_colors = [C_GREEN, C_SAND, C_TERRA]
    for i, (start, end, col) in enumerate([(np.pi, 2*np.pi/3, zone_colors[0]),
                                            (2*np.pi/3, np.pi/3, zone_colors[1]),
                                            (np.pi/3, 0, zone_colors[2])]):
        t = np.linspace(start, end, 100)
        ax.fill_between(t, 0.6, 1.0, color=col, alpha=0.55)

    # Aiguille
    needle_angle = np.pi * (1 - score)
    ax.annotate("", xy=(needle_angle, 0.85), xytext=(0, 0),
                arrowprops=dict(arrowstyle="-|>", color=C_INK, lw=2.5))

    ax.set_ylim(0, 1.1)
    ax.set_theta_zero_location("W")
    ax.set_theta_direction(-1)
    ax.set_thetamin(0)
    ax.set_thetamax(180)
    ax.set_rticks([])
    ax.set_xticks([0, np.pi/4, np.pi/2, 3*np.pi/4, np.pi])
    ax.set_xticklabels(["100%", "75%", "50%", "25%", "0%"], fontsize=9)
    ax.grid(False)
    ax.spines["polar"].set_visible(False)
    ax.set_title(f"Gauge — Indice de menace : {score*100:.0f}%", pad=20, fontweight="bold")
    return fig_to_bytes(fig)


# ---------------------------------------------------------------------------
# 13. Ridgeline plot (comparaison multi-armes)
# ---------------------------------------------------------------------------
def chart_ridgeline():
    fig, axes = plt.subplots(3, 1, figsize=(6, 5), sharex=True)
    datasets = [
        (dmg,         "Bolter lourd",   C_TERRA),
        (dmg_weapon2, "Lance-plasma",   C_BLUE),
        (sim_damage() + rng.integers(4, 8, 2000), "Lascannon", C_GREEN),
    ]
    for ax, (data, label, col) in zip(axes, datasets):
        kde = stats.gaussian_kde(data, bw_method=0.35)
        x = np.linspace(0, 30, 300)
        y = kde(x)
        ax.fill_between(x, y, alpha=0.45, color=col)
        ax.plot(x, y, color=col, linewidth=2)
        ax.axvline(np.mean(data), color=C_INK, linestyle="--", linewidth=1.2, alpha=0.7)
        ax.set_ylabel(label, fontsize=9, rotation=0, ha="right", labelpad=6)
        ax.set_yticks([])
        ax.grid(True, axis="x", alpha=0.4)
        ax.spines["top"].set_visible(False)
        ax.spines["right"].set_visible(False)
        ax.spines["left"].set_visible(False)
    axes[-1].set_xlabel("Dégâts")
    fig.suptitle("Ridgeline — Comparaison multi-armes", fontweight="bold", y=1.01)
    fig.tight_layout()
    return fig_to_bytes(fig)


# ---------------------------------------------------------------------------
# 14. Heatmap probabilité (rounds × dégâts seuil)
# ---------------------------------------------------------------------------
def chart_heatmap():
    fig, ax = plt.subplots(figsize=(6, 4))
    p_destroy = np.mean(dmg >= 18)
    rounds = np.arange(1, 7)
    thresholds = np.arange(4, 20, 2)
    matrix = np.zeros((len(thresholds), len(rounds)))
    for i, thr in enumerate(thresholds):
        p = np.mean(dmg >= thr)
        for j, r in enumerate(rounds):
            matrix[i, j] = (1 - (1 - p) ** r) * 100
    im = ax.imshow(matrix, cmap="YlOrRd", aspect="auto", vmin=0, vmax=100)
    plt.colorbar(im, ax=ax, label="Probabilité (%)")
    ax.set_xticks(range(len(rounds)))
    ax.set_xticklabels([f"Round {r}" for r in rounds], fontsize=9)
    ax.set_yticks(range(len(thresholds)))
    ax.set_yticklabels([f"≥ {t} dég." for t in thresholds], fontsize=9)
    for i in range(len(thresholds)):
        for j in range(len(rounds)):
            v = matrix[i, j]
            ax.text(j, i, f"{v:.0f}%", ha="center", va="center",
                    fontsize=8, color="white" if v > 60 else C_INK)
    ax.set_title("Heatmap — P(atteindre X dégâts en N rounds)", pad=8, fontweight="bold")
    return fig_to_bytes(fig)


# ---------------------------------------------------------------------------
# 15. Bar chart avec intervalles de confiance
# ---------------------------------------------------------------------------
def chart_error_bars():
    fig, ax = plt.subplots(figsize=(6, 3.5))
    weapons = ["Bolter lourd", "Lance-plasma", "Lascannon"]
    means_ = [np.mean(dmg), np.mean(dmg_weapon2), np.mean(sim_damage()) + 5.5]
    p10s   = [np.percentile(dmg, 10), np.percentile(dmg_weapon2, 10), 2.5]
    p90s   = [np.percentile(dmg, 90), np.percentile(dmg_weapon2, 90), 16.0]
    yerr_low  = [m - p10 for m, p10 in zip(means_, p10s)]
    yerr_high = [p90 - m for m, p90 in zip(means_, p90s)]
    colors = [C_TERRA, C_BLUE, C_GREEN]
    x = np.arange(len(weapons))
    bars = ax.bar(x, means_, color=colors, edgecolor=C_INK, linewidth=0.7,
                  width=0.5, alpha=0.8)
    ax.errorbar(x, means_, yerr=[yerr_low, yerr_high], fmt="none",
                color=C_INK, linewidth=2, capsize=6, capthick=1.5)
    ax.set_xticks(x)
    ax.set_xticklabels(weapons, fontsize=10)
    for bar, m in zip(bars, means_):
        ax.text(bar.get_x() + bar.get_width() / 2, m + 0.3, f"{m:.1f}",
                ha="center", fontsize=9, fontweight="bold")
    styled_ax(ax, "Bar chart + IC80% — Comparaison armes", "", "Dégâts moyens")
    return fig_to_bytes(fig)


# ---------------------------------------------------------------------------
# Construction du document Word
# ---------------------------------------------------------------------------
def set_font(run, size_pt=11, bold=False, color=None, italic=False):
    run.bold = bold
    run.italic = italic
    run.font.size = Pt(size_pt)
    if color:
        r, g, b = int(color[1:3], 16), int(color[3:5], 16), int(color[5:7], 16)
        run.font.color.rgb = RGBColor(r, g, b)


def add_heading(doc, text, level=1):
    p = doc.add_heading(text, level=level)
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    return p


def add_para(doc, text, italic=False, size=11):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.italic = italic
    run.font.size = Pt(size)
    return p


def add_chart(doc, img_bytes, width_in=5.5, caption=""):
    doc.add_picture(img_bytes, width=Inches(width_in))
    if caption:
        p = doc.add_paragraph(caption)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.runs[0].italic = True
        p.runs[0].font.size = Pt(9)
        p.runs[0].font.color.rgb = RGBColor(0x5A, 0x5A, 0x7A)


def add_table_row(table, cells):
    row = table.add_row()
    for i, val in enumerate(cells):
        row.cells[i].text = val
    return row


def build_doc():
    doc = Document()

    # Marges
    from docx.oxml.ns import qn
    for section in doc.sections:
        section.left_margin   = Inches(1.1)
        section.right_margin  = Inches(1.1)
        section.top_margin    = Inches(1.0)
        section.bottom_margin = Inches(1.0)

    # -----------------------------------------------------------------------
    # PAGE DE TITRE
    # -----------------------------------------------------------------------
    doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("WARHAMMER 40K SIMULATOR")
    set_font(r, 22, bold=True, color=C_INK)

    p2 = doc.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r2 = p2.add_run("Rapport — Représentations statistiques\npour données de probabilité de combat")
    set_font(r2, 14, italic=True, color="#5A5A7A")

    p3 = doc.add_paragraph()
    p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r3 = p3.add_run("Juin 2026  ·  Brainstorm visuel")
    set_font(r3, 10, color="#8A8A9A")

    doc.add_page_break()

    # -----------------------------------------------------------------------
    # INTRODUCTION
    # -----------------------------------------------------------------------
    add_heading(doc, "Introduction", 1)
    add_para(doc,
        "Ce rapport recense les principales familles de représentation graphique "
        "adaptées à des données de probabilité discrètes, telles que celles produites "
        "par le simulateur Monte Carlo. Les exemples sont générés avec des données "
        "synthétiques réalistes (distribution WH40K, ~2000 runs). "
        "Chaque graphique est accompagné d'une description, d'une évaluation et de "
        "recommandations d'usage pour l'application.")
    add_para(doc,
        "Données simulées : distribution de dégâts sur 2000 rounds, mélange "
        "de 18% de rounds à 0 dégât et d'une Poisson(λ=5.5), clippée à 24.",
        italic=True, size=9)

    doc.add_paragraph()

    # -----------------------------------------------------------------------
    # SECTION 1 — DISTRIBUTION INDIVIDUELLE
    # -----------------------------------------------------------------------
    add_heading(doc, "1. Famille — Distribution individuelle", 1)
    add_para(doc,
        "Ces graphiques montrent la forme complète de la distribution sur "
        "l'ensemble des simulations. Ils répondent à la question : "
        "\"Quels dégâts sont les plus probables, et à quel point la distribution est-elle étalée ?\"")

    # 1.1 Histogramme
    add_heading(doc, "1.1 Histogramme", 2)
    add_para(doc,
        "Le graphique actuel de l'app. Chaque barre représente la fréquence d'un résultat "
        "discret précis. Lisible, familier, mais ne donne aucune information sur la densité "
        "relative ou la forme globale.")
    add_para(doc, "Forces : intuitif, lecture directe des fréquences brutes.", italic=True)
    add_para(doc, "Limites : bruité sur les queues, pas de lecture probabiliste directe.", italic=True)
    add_chart(doc, chart_histogram(), caption="Fig. 1.1 — Histogramme classique (actuel dans l'app)")

    # 1.2 KDE
    add_heading(doc, "1.2 KDE — Kernel Density Estimate", 2)
    add_para(doc,
        "Version lissée de l'histogramme. La courbe représente une estimation "
        "continue de la densité de probabilité. Idéale pour percevoir la forme "
        "globale (bimodalité, asymétrie, queues lourdes).")
    add_para(doc, "Forces : élégant, fait ressortir la structure de la distribution.", italic=True)
    add_para(doc, "Limites : peut sur-lisser et masquer des pics réels sur données discrètes.", italic=True)
    add_chart(doc, chart_kde(), caption="Fig. 1.2 — KDE (densité lissée)")

    # 1.3 Hist + KDE
    add_heading(doc, "1.3 Histogramme + KDE superposés", 2)
    add_para(doc,
        "Combinaison des deux : l'histogramme donne les fréquences réelles, "
        "le KDE donne la forme globale. C'est souvent le meilleur compromis "
        "pour un rapport.")
    add_para(doc, "Recommandation : remplacer l'histogramme actuel par cette version.", italic=True)
    add_chart(doc, chart_hist_kde(), caption="Fig. 1.3 — Histogramme + KDE superposés (recommandé)")

    # 1.4 Strip plot
    add_heading(doc, "1.4 Strip plot (jitter plot)", 2)
    add_para(doc,
        "Chaque point représente un run individuel. Le jitter vertical évite "
        "la superposition. Très utile pour les petits échantillons ou pour "
        "montrer la densité concrète des données.")
    add_para(doc, "Forces : montre les données brutes, aucune abstraction.", italic=True)
    add_para(doc, "Limites : illisible au-delà de 500 points si non échantillonné.", italic=True)
    add_chart(doc, chart_strip(), caption="Fig. 1.4 — Strip plot (300 points échantillonnés sur 2000)")

    # 1.5 Raincloud
    add_heading(doc, "1.5 Raincloud plot", 2)
    add_para(doc,
        "Combinaison de trois niveaux de lecture : le violin (cloud) pour la densité, "
        "le box plot pour les statistiques clés, et les points individuels (pluie). "
        "Permet une lecture complète en un seul visuel.")
    add_para(doc, "Forces : le graphique le plus informatif disponible pour une seule distribution.", italic=True)
    add_para(doc, "Limites : demande un peu de lecture pour les non-initiés.", italic=True)
    add_chart(doc, chart_raincloud(), caption="Fig. 1.5 — Raincloud plot (cloud + box + points)")

    doc.add_page_break()

    # -----------------------------------------------------------------------
    # SECTION 2 — RÉSUMÉ STATISTIQUE
    # -----------------------------------------------------------------------
    add_heading(doc, "2. Famille — Résumé statistique", 1)
    add_para(doc,
        "Ces graphiques condensent la distribution en quelques statistiques clés. "
        "Ils répondent à : \"En chiffres, quelles sont les valeurs importantes ?\"")

    # 2.1 Box plot
    add_heading(doc, "2.1 Box plot", 2)
    add_para(doc,
        "Représente le résumé en 5 nombres : minimum, P25, médiane (P50), P75, maximum. "
        "La boîte couvre l'intervalle interquartile (IQR). Les valeurs aberrantes "
        "sont affichées individuellement.")
    add_para(doc, "Forces : standard académique, compact, facilement comparable entre distributions.", italic=True)
    add_para(doc, "Limites : masque complètement la forme de la distribution (bimodalité invisible).", italic=True)
    add_chart(doc, chart_boxplot(), caption="Fig. 2.1 — Box plot horizontal avec annotations de percentiles")

    # 2.2 Violin
    add_heading(doc, "2.2 Violin plot", 2)
    add_para(doc,
        "Extension du box plot : superpose un KDE miroir sur les deux côtés. "
        "La largeur encode la probabilité à chaque valeur. Combine résumé et forme.")
    add_para(doc, "Forces : révèle la bimodalité et les asymétries que le box plot masque.", italic=True)
    add_para(doc, "Limites : légèrement moins intuitif pour un public non technique.", italic=True)
    add_chart(doc, chart_violin(), caption="Fig. 2.2 — Violin plot (distribution + densité)")

    # 2.3 Bar + IC
    add_heading(doc, "2.3 Bar chart avec intervalles de confiance", 2)
    add_para(doc,
        "Idéal pour la comparaison multi-armes. Les barres montrent la moyenne, "
        "les barres d'erreur représentent l'intervalle P10–P90 (80% de confiance). "
        "Donne immédiatement le \"pire cas réaliste\" et le \"meilleur cas réaliste\".")
    add_para(doc, "Forces : parfait pour comparer plusieurs armes côte à côte.", italic=True)
    add_para(doc, "Recommandation : ajouter en mode multi-armes.", italic=True)
    add_chart(doc, chart_error_bars(), caption="Fig. 2.3 — Bar chart avec intervalles P10–P90")

    doc.add_page_break()

    # -----------------------------------------------------------------------
    # SECTION 3 — PROBABILITÉ CUMULATIVE
    # -----------------------------------------------------------------------
    add_heading(doc, "3. Famille — Probabilité cumulative", 1)
    add_para(doc,
        "Ces graphiques répondent à : \"Quelle est la probabilité de faire au moins X dégâts ?\" "
        "C'est la question tactique centrale dans WH40K.")

    # 3.1 CDF
    add_heading(doc, "3.1 CDF — Fonction de répartition", 2)
    add_para(doc,
        "P(dégâts ≤ X). Lecture : pour X = 8, lire la probabilité d'obtenir "
        "8 dégâts ou moins. Utile pour les joueurs qui cherchent à connaître "
        "le \"plancher\" probable.")
    add_para(doc, "Forces : standard statistique, monotone croissante, facile à lire.", italic=True)
    add_chart(doc, chart_cdf(), caption="Fig. 3.1 — CDF (actuellement dans l'app)")

    # 3.2 Survival function
    add_heading(doc, "3.2 Survival function — P(X ≥ x)", 2)
    add_para(doc,
        "L'inverse de la CDF : P(dégâts ≥ X). C'est le graphique le plus "
        "naturel pour WH40K — le joueur veut savoir \"j'ai X% de chances de "
        "faire au moins Y dégâts\". La courbe actuelle de l'app est déjà une "
        "survival function, mais sans annotation des percentiles.")
    add_para(doc, "Recommandation : annoter la médiane et les percentiles P25/P75 sur ce graphique.", italic=True)
    add_chart(doc, chart_survival(), caption="Fig. 3.2 — Survival function avec médiane annotée")

    # 3.3 Bandes de percentiles
    add_heading(doc, "3.3 Bandes de percentiles (confidence bands)", 2)
    add_para(doc,
        "Enrichit la survival function avec des bandes de couleur progressive : "
        "P25–P75 (IQR), P10–P90, et la courbe médiane. Permet de voir "
        "d'un coup d'œil la variabilité et l'intervalle de confiance.")
    add_para(doc, "Recommandation forte : remplacer le graphique cumulatif actuel par cette version.", italic=True)
    add_chart(doc, chart_percentile_bands(),
              caption="Fig. 3.3 — Survival function avec bandes P25/P75 et P10/P90")

    doc.add_page_break()

    # -----------------------------------------------------------------------
    # SECTION 4 — PROCESSUS / SÉQUENCE
    # -----------------------------------------------------------------------
    add_heading(doc, "4. Famille — Processus et séquence", 1)
    add_para(doc,
        "Ces graphiques montrent comment les valeurs se transforment d'une "
        "étape à l'autre. Ils répondent à : \"Où les dés perdent-ils leur efficacité ?\"")

    # 4.1 Funnel
    add_heading(doc, "4.1 Funnel chart", 2)
    add_para(doc,
        "Le graphique le plus naturel pour WH40K. Visualise l'entonnoir "
        "Attaques → Touches → Blessures → Sauvegardées → Dégâts. "
        "Montre immédiatement où la majorité des dés sont perdus (typiquement "
        "sauvegarde ou phase de blessure). Les pourcentages annotés donnent "
        "le rendement de chaque phase.")
    add_para(doc, "Forces : intuitif, actionnable, unique à WH40K dans ce contexte.", italic=True)
    add_para(doc, "Recommandation forte : ajouter en remplacement ou en complément de l'expander actuel.", italic=True)
    add_chart(doc, chart_funnel(), caption="Fig. 4.1 — Funnel chart des phases de combat (valeurs moyennes)")

    doc.add_page_break()

    # -----------------------------------------------------------------------
    # SECTION 5 — SYNTHÈSE / KPI
    # -----------------------------------------------------------------------
    add_heading(doc, "5. Famille — Synthèse et indicateurs", 1)
    add_para(doc,
        "Ces représentations condensent l'information en un signal unique. "
        "Elles répondent à : \"En un mot, c'est bon ou pas ?\"")

    # 5.1 Gauge
    add_heading(doc, "5.1 Gauge — Jauge de menace", 2)
    add_para(doc,
        "Un indicateur synthétique 0–100% calculé depuis les statistiques clés "
        "(taux de destruction × 0.4 + dégâts moyens/PV total × 0.4 + "
        "destruction complète × 0.2). Permet une lecture immédiate du niveau "
        "de danger pour le défenseur.")
    add_para(doc, "Forces : lecture instantanée, très visuel, \"wow effect\".", italic=True)
    add_para(doc, "Limites : l'agrégation est subjective et peut tromper. À utiliser comme indication.", italic=True)
    add_chart(doc, chart_gauge(), caption="Fig. 5.1 — Gauge de menace (score synthétique)")

    doc.add_page_break()

    # -----------------------------------------------------------------------
    # SECTION 6 — COMPARAISON MULTI-ARMES
    # -----------------------------------------------------------------------
    add_heading(doc, "6. Famille — Comparaison multi-distributions", 1)
    add_para(doc,
        "Ces graphiques permettent de comparer plusieurs armes ou configurations "
        "dans un seul visuel.")

    # 6.1 Ridgeline
    add_heading(doc, "6.1 Ridgeline plot", 2)
    add_para(doc,
        "Superpose plusieurs KDE sur des axes décalés verticalement. "
        "Permet de comparer la forme et le centre de plusieurs distributions "
        "en un coup d'œil. Très adapté quand on a 3+ armes à comparer.")
    add_para(doc, "Forces : comparaison de forme, élégant, évite les graphiques multiples.", italic=True)
    add_para(doc, "Limites : les queues peuvent se chevaucher sur des données très étalées.", italic=True)
    add_chart(doc, chart_ridgeline(), caption="Fig. 6.1 — Ridgeline plot (comparaison 3 armes)")

    # 6.2 Heatmap
    add_heading(doc, "6.2 Heatmap — Probabilité rounds × seuils", 2)
    add_para(doc,
        "Matrice de probabilité : en lignes, les seuils de dégâts visés ; "
        "en colonnes, le nombre de rounds. Chaque cellule = probabilité d'avoir "
        "atteint au moins X dégâts cumulés après N rounds. "
        "Répond à : \"En combien de rounds j'ai 80% de chances de tuer l'unité ?\"")
    add_para(doc, "Forces : très dense en information, lecture tactique directe.", italic=True)
    add_para(doc, "Limites : suppose des rounds indépendants (approximation).", italic=True)
    add_chart(doc, chart_heatmap(), caption="Fig. 6.2 — Heatmap probabilité (rounds × seuils de dégâts)")

    doc.add_page_break()

    # -----------------------------------------------------------------------
    # SECTION 7 — REPRÉSENTATIONS NON GRAPHIQUES
    # -----------------------------------------------------------------------
    add_heading(doc, "7. Représentations non graphiques", 1)

    add_heading(doc, "7.1 Phrase narrative générée automatiquement", 2)
    add_para(doc,
        "Une ligne de texte générée depuis les stats de la simulation. "
        "Zéro complexité technique, impact immédiat.")
    p = doc.add_paragraph()
    r = p.add_run(
        "→ Exemple : \"73% de chances de tuer au moins 3 figurines. "
        "Destruction totale dans 28% des cas. "
        "Dans 80% des runs, les dégâts sont compris entre 2 et 11 (P10–P90).\"")
    set_font(r, 11, bold=True, color=C_TERRA)

    add_heading(doc, "7.2 Tableau de percentiles", 2)
    add_para(doc,
        "Un tableau compact résumant les seuils clés en termes de probabilité. "
        "Très utile pour les joueurs qui veulent planifier précisément.")

    table = doc.add_table(rows=1, cols=3)
    table.style = "Light Grid Accent 1"
    hdr = table.rows[0].cells
    hdr[0].text = "Percentile"
    hdr[1].text = "Dégâts"
    hdr[2].text = "Lecture"
    for pct, label in [(10, "Pire cas réaliste"), (25, "Quart défavorable"),
                       (50, "Médiane"), (75, "Quart favorable"), (90, "Meilleur cas réaliste")]:
        val = int(np.percentile(dmg, pct))
        row = table.add_row()
        row.cells[0].text = f"P{pct}"
        row.cells[1].text = str(val)
        row.cells[2].text = label

    doc.add_paragraph()

    add_heading(doc, "7.3 Sparkline de distribution inline", 2)
    add_para(doc,
        "Miniature de l'histogramme affichée dans la metric card elle-même. "
        "Permet de voir la distribution sans scrolling. "
        "Nécessite un composant HTML custom ou une image base64 inline dans Streamlit.")

    doc.add_page_break()

    # -----------------------------------------------------------------------
    # SECTION 8 — SYNTHÈSE ET RECOMMANDATIONS
    # -----------------------------------------------------------------------
    add_heading(doc, "8. Synthèse et recommandations prioritaires", 1)
    add_para(doc,
        "Classement par ratio impact/complexité d'implémentation, "
        "adapté au contexte Streamlit + Plotly.")

    table2 = doc.add_table(rows=1, cols=4)
    table2.style = "Light Grid Accent 1"
    hdr2 = table2.rows[0].cells
    hdr2[0].text = "Priorité"
    hdr2[1].text = "Représentation"
    hdr2[2].text = "Impact"
    hdr2[3].text = "Complexité"

    recs = [
        ("1", "Phrase narrative auto", "★★★★★", "★☆☆☆☆"),
        ("2", "Bandes de percentiles P10/P90 sur survival fn", "★★★★☆", "★★☆☆☆"),
        ("3", "Funnel chart phases de combat", "★★★★★", "★★☆☆☆"),
        ("4", "Histogramme + KDE superposés", "★★★☆☆", "★☆☆☆☆"),
        ("5", "Gauge de menace (go.Indicator Plotly)", "★★★★☆", "★★★☆☆"),
        ("6", "Ridgeline plot (multi-armes)", "★★★★☆", "★★★☆☆"),
        ("7", "Heatmap rounds × seuils", "★★★☆☆", "★★★☆☆"),
        ("8", "Violin / Raincloud plot", "★★★☆☆", "★★★★☆"),
    ]
    for rec in recs:
        row = table2.add_row()
        for i, val in enumerate(rec):
            row.cells[i].text = val

    doc.add_paragraph()
    add_para(doc,
        "Note : toutes les représentations sont faisables avec Plotly (déjà installé). "
        "Le funnel chart et le gauge utilisent go.Funnel et go.Indicator nativement. "
        "Le ridgeline se construit avec des go.Violin empilés. "
        "La phrase narrative est du Python pur (numpy percentiles → f-string).",
        italic=True, size=9)

    # -----------------------------------------------------------------------
    # FOOTER
    # -----------------------------------------------------------------------
    doc.add_page_break()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("Powered by Wahapedia · wahapedia.ru\n"
                  "Warhammer 40,000 © Games Workshop Ltd — fan project, non affilié")
    set_font(r, 9, italic=True, color="#8A8A9A")

    return doc


# ---------------------------------------------------------------------------
# Main
# ---------------------------------------------------------------------------
if __name__ == "__main__":
    print("Génération des graphiques...")
    doc = build_doc()
    out = "rapport_stats_proba.docx"
    doc.save(out)
    print(f"Document sauvegardé : {out}")
